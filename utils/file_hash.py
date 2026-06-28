"""
HashGuard - File Hash Utilities
Path: utils/file_hash.py
Purpose: Compute and compare SHA-256 digests for uploaded files to support
         integrity verification and duplicate detection across all supported formats.
"""

import hashlib
import io
import re
import zipfile
from PIL import Image


def compute_sha256_from_path(file_path, chunk_size=8192):
    """
    Read a file from disk in chunks and return its SHA-256 hex digest.
    """
    digest = hashlib.sha256()
    with open(file_path, "rb") as file_handle:
        while True:
            chunk = file_handle.read(chunk_size)
            if not chunk:
                break
            digest.update(chunk)
    return digest.hexdigest()


def compute_sha256_from_stream(file_stream, chunk_size=8192):
    """
    Compute SHA-256 from an uploaded file stream without persisting it first.
    Resets the stream position to the beginning after hashing.
    """
    file_stream.seek(0)
    digest = hashlib.sha256()
    while True:
        chunk = file_stream.read(chunk_size)
        if not chunk:
            break
        digest.update(chunk)
    file_stream.seek(0)
    return digest.hexdigest()


def compute_content_hash_from_stream(file_stream, extension):
    """
    Compute a content-based hash that ignores metadata, timestamps, and filenames.
    Supports PDF, TXT, DOCX, XLSX, PPTX, PNG, JPG, JPEG, and ZIP files.
    Returns the SHA-256 hex digest of the normalized content payload.
    """
    file_stream.seek(0)
    ext = (extension or "").lower().lstrip(".")
    content_payload = ""

    try:
        if ext == "txt":
            content_payload = _extract_txt(file_stream)
        elif ext == "docx":
            content_payload = _extract_docx(file_stream)
        elif ext == "pdf":
            content_payload = _extract_pdf(file_stream)
        elif ext == "zip":
            content_payload = _extract_zip(file_stream)
        elif ext == "xlsx":
            content_payload = _extract_xlsx(file_stream)
        elif ext == "pptx":
            content_payload = _extract_pptx(file_stream)
        elif ext in ("png", "jpg", "jpeg"):
            content_payload = _extract_image(file_stream)
        else:
            content_payload = _extract_fallback(file_stream)
    except Exception as e:
        print(f"[DEBUG content_hash error] ext={ext} err={e}")
        file_stream.seek(0)
        content_payload = compute_sha256_from_stream(file_stream)

    file_stream.seek(0)
    normalized = _normalize_text(content_payload)
    final_raw = f"type:{ext};content:{normalized}"
    content_hash = hashlib.sha256(final_raw.encode("utf-8")).hexdigest()
    print(f"[DEBUG content_hash] ext={ext} norm_len={len(normalized)} hash={content_hash}")
    return content_hash


def _extract_txt(file_stream):
    try:
        file_stream.seek(0)
        return file_stream.read().decode("utf-8", errors="replace")
    except Exception:
        return ""


def _extract_docx(file_stream):
    try:
        from docx import Document
        file_stream.seek(0)
        document = Document(file_stream)
        text_parts = []
        for p in document.paragraphs:
            if p.text:
                text_parts.append(p.text)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text_parts.append(cell.text)
        return "\n".join(text_parts)
    except Exception:
        return _extract_openxml_text(file_stream, r"word/document\.xml")


def _extract_pdf(file_stream):
    try:
        from pypdf import PdfReader
        file_stream.seek(0)
        reader = PdfReader(file_stream)
        pages = []
        for page in reader.pages:
            text = page.extract_text() or ""
            pages.append(text)
        return "\n".join(pages)
    except Exception:
        return _extract_fallback(file_stream)


def _extract_zip(file_stream):
    try:
        file_stream.seek(0)
        with zipfile.ZipFile(file_stream, "r") as zf:
            file_entries = []
            for name in sorted(zf.namelist()):
                if name.endswith("/"):
                    continue
                try:
                    data = zf.read(name)
                    data_hash = hashlib.sha256(data).hexdigest()
                    file_entries.append(f"{name}:{data_hash}")
                except Exception:
                    continue
            return "\n".join(file_entries)
    except Exception:
        return _extract_fallback(file_stream)


def _extract_xlsx(file_stream):
    try:
        return _extract_openxml_text(file_stream, r"xl/(worksheets/sheet|sharedStrings)\.xml")
    except Exception:
        return _extract_fallback(file_stream)


def _extract_pptx(file_stream):
    try:
        return _extract_openxml_text(file_stream, r"ppt/slides/slide.*\.xml")
    except Exception:
        return _extract_fallback(file_stream)


def _extract_openxml_text(file_stream, file_pattern):
    file_stream.seek(0)
    extracted_texts = []
    pattern = re.compile(file_pattern)
    with zipfile.ZipFile(file_stream, "r") as zf:
        for name in sorted(zf.namelist()):
            if pattern.search(name):
                xml_content = zf.read(name).decode("utf-8", errors="replace")
                text_content = re.sub(r"<[^>]+>", " ", xml_content)
                extracted_texts.append(text_content)
    return " ".join(extracted_texts)


def _extract_image(file_stream):
    try:
        file_stream.seek(0)
        with Image.open(file_stream) as img:
            img_rgba = img.convert("RGBA")
            pixel_bytes = img_rgba.tobytes()
            return f"dim:{img.size[0]}x{img.size[1]};pixels:{hashlib.sha256(pixel_bytes).hexdigest()}"
    except Exception:
        file_stream.seek(0)
        return compute_sha256_from_stream(file_stream)


def _extract_fallback(file_stream):
    file_stream.seek(0)
    return compute_sha256_from_stream(file_stream)


def _normalize_text(text):
    if not text:
        return ""
    normalized = " ".join(text.lower().split())
    return normalized
